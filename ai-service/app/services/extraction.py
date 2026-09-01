from io import BytesIO
from typing import Optional, Tuple

import fitz
import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader

from app.config import SUPPORTED_CONTENT_TYPES


class UnsupportedContentTypeError(ValueError):
    pass


class TextExtractionError(ValueError):
    pass


def extract_text(data: bytes, content_type: Optional[str]) -> Tuple[str, str]:
    normalized_content_type = normalize_content_type(content_type)
    if normalized_content_type not in SUPPORTED_CONTENT_TYPES:
        raise UnsupportedContentTypeError("Unsupported content type")

    try:
        if normalized_content_type == "application/pdf":
            text = extract_pdf_text(data)
        elif normalized_content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_docx_text(data)
        elif normalized_content_type in {"image/jpeg", "image/png"}:
            text = extract_image_text(data)
        else:
            text = data.decode("utf-8")
    except Exception as exc:
        raise TextExtractionError("Unable to extract text") from exc

    return normalized_content_type, text


def normalize_content_type(content_type: Optional[str]) -> str:
    return (content_type or "").split(";", 1)[0].strip().lower()


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    embedded_text = "\n".join(page_text for page_text in pages if page_text)
    if not needs_scanned_pdf_ocr(pages):
        return embedded_text

    return extract_scanned_pdf_text(data)


def needs_scanned_pdf_ocr(page_texts: list[str]) -> bool:
    """Only OCR when the embedded layer is absent, preserving every text PDF's fast path."""
    embedded_text = "\n".join(page_texts)
    significant_characters = sum(character.isalnum() for character in embedded_text)
    return significant_characters == 0


def extract_scanned_pdf_text(data: bytes) -> str:
    # PyMuPDF renders pages directly, avoiding pdf2image's Poppler system dependency.
    document = fitz.open(stream=data, filetype="pdf")
    try:
        page_texts = []
        for page in document:
            pixmap = page.get_pixmap(dpi=200, alpha=False)
            with Image.open(BytesIO(pixmap.tobytes("png"))) as image:
                page_texts.append(pytesseract.image_to_string(image))
        return "\n".join(page_text for page_text in page_texts if page_text)
    finally:
        document.close()


def extract_image_text(data: bytes) -> str:
    with Image.open(BytesIO(data)) as image:
        return pytesseract.image_to_string(image)


def extract_docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    sections = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                sections.append("\t".join(cells))
    return "\n".join(sections)
